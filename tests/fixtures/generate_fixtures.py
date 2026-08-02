"""Regenerate the repository's small synthetic document fixtures."""

from pathlib import Path
from zipfile import ZipFile, ZipInfo

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

FIXTURE_ROOT = Path(__file__).parent
PDF_ROOT = FIXTURE_ROOT / "pdf"
DOCX_ROOT = FIXTURE_ROOT / "docx"


def _add_text_page(writer: PdfWriter, lines: list[str]) -> None:
    page = writer.add_blank_page(width=612, height=792)
    if not lines:
        return

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )

    commands = ["BT", "/F1 12 Tf", "72 720 Td"]
    for index, line in enumerate(lines):
        if index:
            commands.append("0 -24 Td")
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        commands.append(f"({escaped}) Tj")
    commands.append("ET")

    stream = DecodedStreamObject()
    stream.set_data("\n".join(commands).encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)


def _write_pdf(name: str, pages: list[list[str]], *, encrypted: bool = False) -> None:
    writer = PdfWriter()
    for lines in pages:
        _add_text_page(writer, lines)
    if encrypted:
        writer.encrypt("synthetic-test-password")
    with (PDF_ROOT / name).open("wb") as fixture:
        writer.write(fixture)


def _write_non_text_pdf() -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    stream = DecodedStreamObject()
    stream.set_data(b"72 720 72 24 re S")
    page[NameObject("/Contents")] = writer._add_object(stream)
    with (PDF_ROOT / "non-extractable.pdf").open("wb") as fixture:
        writer.write(fixture)


def _write_docx() -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Excluded Header"
    document.sections[0].footer.paragraphs[0].text = "Excluded Footer"
    document.add_paragraph("Before table")
    first_table = document.add_table(rows=2, cols=2)
    first_table.cell(0, 0).text = "A1"
    first_table.cell(0, 1).text = "B1"
    document.add_paragraph("")
    document.add_paragraph("Between tables")
    second_table = document.add_table(rows=1, cols=3)
    second_table.cell(0, 0).text = "Left"
    second_table.cell(0, 2).text = "Right"
    document.add_paragraph("After table")
    path = DOCX_ROOT / "ordered-content.docx"
    document.save(path)
    _normalize_docx_archive(path)


def _normalize_docx_archive(path: Path) -> None:
    normalized_path = path.with_suffix(".normalized.docx")
    with ZipFile(path, "r") as source, ZipFile(normalized_path, "w") as target:
        for original in sorted(source.infolist(), key=lambda entry: entry.filename):
            entry = ZipInfo(original.filename, date_time=(1980, 1, 1, 0, 0, 0))
            entry.compress_type = original.compress_type
            entry.create_system = original.create_system
            entry.external_attr = original.external_attr
            entry.internal_attr = original.internal_attr
            target.writestr(entry, source.read(original.filename))
    normalized_path.replace(path)


def main() -> None:
    PDF_ROOT.mkdir(parents=True, exist_ok=True)
    DOCX_ROOT.mkdir(parents=True, exist_ok=True)
    _write_pdf(
        "multi-page.pdf",
        [["First page text"], [], ["Third page text"]],
    )
    _write_pdf(
        "repeated-margins.pdf",
        [
            [
                "Shared Header",
                "Page one body",
                "Shared Header",
                "Additional one",
                "Page 1 of 3",
            ],
            [
                "Shared Header",
                "Page two body",
                "Another detail",
                "Additional two",
                "Page 2 of 3",
            ],
            [
                "Shared Header",
                "Page three body",
                "Final detail",
                "Additional three",
                "Page 3 of 3",
            ],
        ],
    )
    _write_pdf(
        "single-page.pdf",
        [["Only Page Heading", "Important body", "Page 1"]],
    )
    _write_pdf("empty.pdf", [[]])
    _write_pdf("encrypted.pdf", [["Secret text"]], encrypted=True)
    _write_non_text_pdf()
    (PDF_ROOT / "malformed.pdf").write_bytes(b"This is not a PDF file.\n")
    _write_docx()
    (DOCX_ROOT / "malformed.docx").write_bytes(b"This is not a DOCX package.\n")


if __name__ == "__main__":
    main()
