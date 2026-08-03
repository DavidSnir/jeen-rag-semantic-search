"""Tests for ordered DOCX body extraction."""

from pathlib import Path

import pytest
from docx import Document

from rag_app.exceptions import DocumentExtractionError, EmptyDocumentError
from rag_app.extractors import extract_document

DOCX_FIXTURES = Path(__file__).parents[1] / "fixtures" / "docx"


def test_docx_extracts_paragraphs_and_table_rows_in_body_order() -> None:
    document = extract_document(DOCX_FIXTURES / "ordered-content.docx")

    assert len(document.units) == 5
    assert [unit.text for unit in document.units] == [
        "Before table",
        "A1 | B1",
        "Between tables",
        "Left | | Right",
        "After table",
    ]
    assert [unit.position for unit in document.units] == list(range(5))


def test_docx_excludes_section_headers_and_footers() -> None:
    document = extract_document(DOCX_FIXTURES / "ordered-content.docx")
    combined = "\n".join(unit.text for unit in document.units)

    assert "Excluded Header" not in combined
    assert "Excluded Footer" not in combined


def test_docx_units_have_null_page_numbers_and_canonical_metadata() -> None:
    document = extract_document(DOCX_FIXTURES / "ordered-content.docx")

    assert document.source_file == "ordered-content.docx"
    assert document.source_type == "DOCX"
    assert all(unit.page_number is None for unit in document.units)
    assert str(DOCX_FIXTURES) not in repr(document)


def test_docx_preserves_unicode_and_punctuation(tmp_path: Path) -> None:
    path = tmp_path / "unicode.docx"
    source = Document()
    source.add_paragraph("Résumé: π ≈ 3.14 — 日本語!")
    source.save(path)

    document = extract_document(path)

    assert document.units[0].text == "Résumé: π ≈ 3.14 — 日本語!"


def test_docx_skips_empty_paragraphs_rows_and_cells(tmp_path: Path) -> None:
    path = tmp_path / "empty-elements.docx"
    source = Document()
    source.add_paragraph("  ")
    empty_table = source.add_table(rows=1, cols=2)
    empty_table.cell(0, 0).text = "\t"
    source.add_paragraph("Meaningful")
    partial_table = source.add_table(rows=1, cols=3)
    partial_table.cell(0, 1).text = "center"
    source.save(path)

    document = extract_document(path)

    assert [unit.text for unit in document.units] == ["Meaningful", "| center |"]


def test_docx_preserves_nested_table_text_inside_a_cell(tmp_path: Path) -> None:
    path = tmp_path / "nested-table.docx"
    source = Document()
    outer = source.add_table(rows=1, cols=1)
    outer.cell(0, 0).paragraphs[0].text = "Outer"
    nested = outer.cell(0, 0).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "Nested left"
    nested.cell(0, 1).text = "Nested right"
    source.save(path)

    document = extract_document(path)

    assert document.units[0].text == "Outer Nested left | Nested right"


def test_empty_docx_is_rejected(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)

    with pytest.raises(EmptyDocumentError) as caught:
        extract_document(path)

    assert "does not contain extractable text" in str(caught.value)
    assert "empty.docx" not in str(caught.value)
    assert "OCR" not in str(caught.value)


def test_docx_with_only_empty_table_cells_is_rejected(tmp_path: Path) -> None:
    path = tmp_path / "empty-table.docx"
    source = Document()
    source.add_table(rows=2, cols=2)
    source.save(path)

    with pytest.raises(EmptyDocumentError):
        extract_document(path)


def test_malformed_docx_is_reported_without_parser_details() -> None:
    path = DOCX_FIXTURES / "malformed.docx"

    with pytest.raises(DocumentExtractionError) as caught:
        extract_document(path)

    assert "malformed.docx" in str(caught.value)
    assert str(DOCX_FIXTURES) not in str(caught.value)
    assert "package not found" not in str(caught.value).lower()
    assert caught.value.__cause__ is not None
